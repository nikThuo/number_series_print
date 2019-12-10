import random
import docx
# import win32ui
# from win32com import client
import time
import subprocess
import os
import tempfile

print(('%06x' % random.randrange(16**6)).upper())
# a = ('%06x' % random.randrange(16**6)).upper()
# print("Length is: ", len(a))
# filename = tempfile.mktemp(".doc")
# open(filename, "w").write(('%06x' % random.randrange(16**6)).upper())
# os.system("lpr -P 192.168.8.34 filename")

# def generateCode():
#     codes = []
#
#     while len(codes) < 100:
#         code = ('%06x' % random.randrange(16**6)).upper()
#         if code in codes:
#             continue
#         else:
#             codes += [code]
#
#     doc = docx.Document()
#     doc.add_heading('Item Code')
#     doc.add_paragraph(code)
#     return doc

    # for c in codes:
    #     print('\n', c)

# c = generateCode()
# print(c)

# os.getcwd()

# generateCode()

# def printCode():
#     c = generateCode()
#     print(c.save())
#     loc = os.path.abspath(str(c))
#     # os.system('start c')
#     os.startfile(os.path.abspath(str(c)))
#
#     dc = win32ui.CreateDC()
#     dc.CreatePrinterDC()
#     dc.StartDoc(str(c))
#
# printCode()

# def printCode():
#     c = generateCode()
#
#     word = client.Dispatch("Word.Application")
#     word.Documents.Open(c)
#     word.ActiveDocument.PrintOut()
#     time.sleep(2)
#     word.ActiveDocument.Close()
#     # file = doc.save('new.docx')
#     #
#     # word = client.Dispatch("Word.Application")
#     # word.Documents.Open(file)
#     # word.ActiveDocument.PrintOut()
#     # time.sleep(2)
#     # word.ActiveDocument.Close()
#
# # word.Quit()
#
# printCode()


# def checkPrinter():
#     conn = cups.Connection()
#     printers = conn.getPrinters()
#
#     for printer in printers:
#         print(printer, printers[printer]["device-uri"])
#
# checkPrinter()




# def printCode():
#     code = generateCode()
#
#     doc = docx.Document()
#
#     doc.add_heading('Item Code')
#     doc.add_paragraph(code)
#
#     # file = doc.save('new.docx')
#
#     # lpr = subprocess.Popen("usr/bin/lpr", stdin=subprocess.PIPE)
#     # lpr.stdin.write(doc)
#
#     # filename = tempfile.mktemp(".doc")
#     # open(filename)
#     #
#     # os.startfile(filename, "print")
#
#     os.system(lpr, doc)
#
#     print(lpr)
#
#
# printCode()